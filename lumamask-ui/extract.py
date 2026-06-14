"""
extract.py — Plain-text extraction from uploaded document files.

Standalone by design: this module imports NO lumamask / presidio / spaCy code,
so its logic can be unit-tested in a lightweight environment without the NLP
stack. The parser libraries (pypdf, python-docx, striprtf) are imported lazily
inside each handler, so a missing optional dependency only breaks the one
format that needs it — every other format still works and the error message
names the package to install.

Public API
----------
    SUPPORTED_EXTENSIONS : frozenset[str]   — e.g. {".txt", ".pdf", ".docx", ...}
    MAX_UPLOAD_BYTES     : int              — hard size cap (per file)
    ExtractionError(Exception)              — raised for any user-facing failure
    extract_text(filename, data) -> str     — dispatch on extension, return text

The returned text is what gets dropped into the document textarea; from there
the existing /api/run pipeline treats it exactly like pasted text.
"""

from __future__ import annotations

import io
import os

# Per-file upload cap. Documents this tool targets (invoices, letters, quotes)
# are tiny; 25 MB is generous while still bounding memory for a local server.
MAX_UPLOAD_BYTES = 25 * 1024 * 1024

# Extensions we read by decoding bytes directly — no third-party parser needed.
# Covers the common "plain text family": notes, markdown, delimited data, logs.
_PLAINTEXT_EXTENSIONS = frozenset(
    {".txt", ".text", ".md", ".markdown", ".csv", ".tsv", ".log", ".rst"}
)

# Extensions that need a dedicated binary parser.
_BINARY_EXTENSIONS = frozenset({".pdf", ".docx", ".rtf"})

SUPPORTED_EXTENSIONS = _PLAINTEXT_EXTENSIONS | _BINARY_EXTENSIONS


class ExtractionError(Exception):
    """A user-facing extraction failure (bad format, missing dep, corrupt file)."""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _ext(filename: str) -> str:
    """Return the lowercased extension of *filename*, including the dot."""
    return os.path.splitext(filename or "")[1].lower()


def _decode_plaintext(data: bytes) -> str:
    """
    Decode raw bytes to text for the plain-text family.

    Tries UTF-8 (with BOM handling) first, then a couple of common fallbacks,
    and finally UTF-8 with replacement so a stray byte never hard-fails an
    otherwise readable document.
    """
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # Last resort: never raise on decode — keep the readable parts.
    return data.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Per-format handlers (parser libs imported lazily, inside each handler)
# ---------------------------------------------------------------------------

def _extract_pdf(data: bytes) -> str:
    try:
        import pypdf
    except ImportError as exc:  # pragma: no cover - exercised via monkeypatch
        raise ExtractionError(
            "Reading PDF files requires the 'pypdf' package (pip install pypdf)."
        ) from exc

    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not open PDF file: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        # Try the common empty-password case; if it stays locked, tell the user.
        try:
            reader.decrypt("")
        except Exception:
            pass
        if getattr(reader, "is_encrypted", False):
            raise ExtractionError(
                "This PDF is password-protected. Remove the password and try again."
            )

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            # One unreadable page shouldn't sink the whole document.
            pages.append("")

    text = "\n".join(pages).strip()
    if not text:
        raise ExtractionError(
            "No selectable text found in this PDF. It may be a scanned image "
            "(OCR is not supported) — paste the text manually instead."
        )
    return text


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - exercised via monkeypatch
        raise ExtractionError(
            "Reading DOCX files requires the 'python-docx' package "
            "(pip install python-docx)."
        ) from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not open DOCX file: {exc}") from exc

    blocks = [p.text for p in document.paragraphs]
    # Tables carry the structured data (amounts, line items) we most care about.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))

    text = "\n".join(blocks).strip()
    if not text:
        raise ExtractionError("No text found in this DOCX file.")
    return text


def _extract_rtf(data: bytes) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError as exc:  # pragma: no cover - exercised via monkeypatch
        raise ExtractionError(
            "Reading RTF files requires the 'striprtf' package "
            "(pip install striprtf)."
        ) from exc

    raw = _decode_plaintext(data)  # RTF is ASCII-ish text with control words
    try:
        text = rtf_to_text(raw).strip()
    except Exception as exc:
        raise ExtractionError(f"Could not parse RTF file: {exc}") from exc

    if not text:
        raise ExtractionError("No text found in this RTF file.")
    return text


_HANDLERS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".rtf": _extract_rtf,
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text(filename: str, data: bytes) -> str:
    """
    Extract plain text from an uploaded file's *data*, dispatching on the
    extension of *filename*.

    Parameters
    ----------
    filename : original filename (only its extension is used).
    data     : raw file bytes.

    Returns
    -------
    The extracted text, stripped of leading/trailing whitespace.

    Raises
    ------
    ExtractionError
        Empty upload, file too large, unsupported extension, missing parser
        dependency, or a corrupt / unreadable / textless document.
    """
    if not data:
        raise ExtractionError("The uploaded file is empty.")
    if len(data) > MAX_UPLOAD_BYTES:
        mb = MAX_UPLOAD_BYTES // (1024 * 1024)
        raise ExtractionError(f"File is too large (limit {mb} MB).")

    ext = _ext(filename)

    if ext in _PLAINTEXT_EXTENSIONS or ext == "":
        # Unknown/extensionless uploads are treated as plain text rather than
        # rejected — the textarea is the safety net and the user sees the result.
        text = _decode_plaintext(data).strip()
        if not text:
            raise ExtractionError("No text found in this file.")
        return text

    handler = _HANDLERS.get(ext)
    if handler is None:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ExtractionError(
            f"Unsupported file type '{ext}'. Supported formats: {supported}."
        )
    return handler(data)
