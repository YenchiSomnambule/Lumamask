"""
test_extract.py — Tests for the standalone text-extraction module.

These tests import only `extract` (no lumamask / presidio / spaCy), so they
run in any environment with pypdf, python-docx and striprtf installed. Binary
fixtures (PDF / DOCX / RTF) are built in-memory, so no binaries live in git.
"""

from __future__ import annotations

import io
import sys

import pytest

import extract
from extract import ExtractionError, extract_text


# ---------------------------------------------------------------------------
# Fixture builders (produce real files in-memory)
# ---------------------------------------------------------------------------

def _make_pdf(text: str) -> bytes:
    """Build a minimal one-page PDF with a single extractable text run."""
    esc = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content = f"BT /F1 24 Tf 72 700 Td ({esc}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n"
        + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_pos = len(pdf)
    pdf += b"xref\n0 " + str(len(objects) + 1).encode() + b"\n"
    pdf += b"0000000000 65535 f \n"
    for off in offsets:
        pdf += f"{off:010d} 00000 n \n".encode()
    pdf += b"trailer\n<< /Size " + str(len(objects) + 1).encode() + b" /Root 1 0 R >>\n"
    pdf += b"startxref\n" + str(xref_pos).encode() + b"\n%%EOF"
    return bytes(pdf)


def _make_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    import docx
    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    if table:
        t = document.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, val in enumerate(row):
                t.rows[r].cells[c].text = val
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


_RTF_SAMPLE = (
    r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}"
    r"\f0\fs24 Hello from RTF, John Smith.\par "
    r"Amount: \$1,234.56\par}"
).encode("latin-1")


# ---------------------------------------------------------------------------
# Module surface
# ---------------------------------------------------------------------------

class TestSupportedExtensions:

    def test_core_formats_present(self):
        for ext in (".txt", ".md", ".csv", ".pdf", ".docx", ".rtf"):
            assert ext in extract.SUPPORTED_EXTENSIONS

    def test_unsupported_format_absent(self):
        assert ".exe" not in extract.SUPPORTED_EXTENSIONS
        assert ".png" not in extract.SUPPORTED_EXTENSIONS


# ---------------------------------------------------------------------------
# Plain-text family
# ---------------------------------------------------------------------------

class TestPlainText:

    def test_txt_utf8(self):
        assert extract_text("notes.txt", "Héllo wörld".encode("utf-8")) == "Héllo wörld"

    def test_strips_surrounding_whitespace(self):
        assert extract_text("a.txt", b"  \n trimmed \n  ") == "trimmed"

    def test_markdown_and_csv_and_log(self):
        assert extract_text("a.md", b"# Title") == "# Title"
        assert extract_text("a.csv", b"name,amount\nAcme,5") == "name,amount\nAcme,5"
        assert extract_text("a.log", b"INFO ok") == "INFO ok"

    def test_utf8_bom_is_stripped(self):
        data = b"\xef\xbb\xbf" + "with bom".encode("utf-8")
        assert extract_text("a.txt", data) == "with bom"

    def test_cp1252_fallback(self):
        # 0xA3 is an invalid UTF-8 lead byte; cp1252 decodes it as £.
        assert extract_text("a.txt", b"Price \xa3100") == "Price £100"

    def test_extensionless_treated_as_text(self):
        assert extract_text("README", b"plain content") == "plain content"

    def test_blank_plaintext_rejected(self):
        with pytest.raises(ExtractionError, match="No text"):
            extract_text("a.txt", b"   \n\t ")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

class TestPdf:

    def test_extracts_text(self):
        out = extract_text("doc.pdf", _make_pdf("Invoice for Acme Corp"))
        assert "Invoice for Acme Corp" in out

    def test_corrupt_pdf_raises(self):
        with pytest.raises(ExtractionError, match="Could not open PDF"):
            extract_text("doc.pdf", b"%PDF-1.4 not really a pdf")

    def test_textless_pdf_reports_scanned_hint(self):
        # A valid PDF whose only content draws no text → "no selectable text".
        blank = _make_pdf("").replace(b"BT /F1 24 Tf 72 700 Td () Tj ET", b" ")
        with pytest.raises(ExtractionError, match="No selectable text"):
            extract_text("scan.pdf", blank)

    def test_missing_pypdf_dependency(self, monkeypatch):
        monkeypatch.setitem(sys.modules, "pypdf", None)
        with pytest.raises(ExtractionError, match="requires the 'pypdf' package"):
            extract_text("doc.pdf", _make_pdf("x"))


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

class TestDocx:

    def test_extracts_paragraphs(self):
        data = _make_docx(["First line.", "Second line, $500.00."])
        out = extract_text("doc.docx", data)
        assert "First line." in out
        assert "Second line, $500.00." in out

    def test_extracts_table_cells(self):
        data = _make_docx(["Header"], table=[["Item", "Total"], ["Widget", "$9.99"]])
        out = extract_text("doc.docx", data)
        assert "Widget" in out
        assert "$9.99" in out

    def test_corrupt_docx_raises(self):
        with pytest.raises(ExtractionError, match="Could not open DOCX"):
            extract_text("doc.docx", b"PK\x03\x04 not a real docx")

    def test_missing_python_docx_dependency(self, monkeypatch):
        data = _make_docx(["x"])  # build the fixture before hiding the module
        monkeypatch.setitem(sys.modules, "docx", None)
        with pytest.raises(ExtractionError, match="requires the 'python-docx' package"):
            extract_text("doc.docx", data)


# ---------------------------------------------------------------------------
# RTF
# ---------------------------------------------------------------------------

class TestRtf:

    def test_extracts_text(self):
        out = extract_text("doc.rtf", _RTF_SAMPLE)
        assert "John Smith" in out
        assert "1,234.56" in out

    def test_missing_striprtf_dependency(self, monkeypatch):
        monkeypatch.setitem(sys.modules, "striprtf", None)
        monkeypatch.setitem(sys.modules, "striprtf.striprtf", None)
        with pytest.raises(ExtractionError, match="requires the 'striprtf' package"):
            extract_text("doc.rtf", _RTF_SAMPLE)


# ---------------------------------------------------------------------------
# Guards: size, empty, unsupported
# ---------------------------------------------------------------------------

class TestGuards:

    def test_empty_upload_rejected(self):
        with pytest.raises(ExtractionError, match="empty"):
            extract_text("a.txt", b"")

    def test_oversize_upload_rejected(self):
        big = b"x" * (extract.MAX_UPLOAD_BYTES + 1)
        with pytest.raises(ExtractionError, match="too large"):
            extract_text("a.txt", big)

    def test_unsupported_extension_lists_supported(self):
        with pytest.raises(ExtractionError, match="Unsupported file type '.xlsx'"):
            extract_text("sheet.xlsx", b"anything")

    def test_extension_case_insensitive(self):
        out = extract_text("DOC.PDF", _make_pdf("Mixed Case Ext"))
        assert "Mixed Case Ext" in out
